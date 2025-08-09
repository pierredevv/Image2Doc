# Lógica de OCR y reconstrucción de texto 
# archivo no necesario (no se usa en el proyecto actual)

import os 
from PIL import Image
import pytesseract
import subprocess
import logging
import tempfile
import re
import cv2
import numpy as np

# Configurar logging
logger = logging.getLogger(__name__)

# Ruta de Tesseract (ajustar según tu sistema)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

from docx import Document
from openpyxl import Workbook

# Función para obtener idiomas disponibles
def get_available_languages():
    try:
        # Ejecutar comando para obtener idiomas instalados
        cmd = [pytesseract.pytesseract.tesseract_cmd, '--list-langs']
        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
        
        # Parsear resultado
        langs = result.stdout.strip().split('\n')[1:]  # Más robusto
        return [lang.strip() for lang in langs if lang.strip()]
    except Exception as e:
        logger.error(f"Error obteniendo idiomas: {str(e)}")
        return ['eng', 'spa']  # Idiomas de respaldo

# Validar formato de idioma
def validate_language(lang):
    available_langs = get_available_languages()
    # Permitir combinaciones como spa+eng
    requested_langs = re.split(r'[+,\s]', lang)  # Más flexible
    
    valid_langs = []
    for l in requested_langs:
        if l and l in available_langs:  # Verificar no vacío
            valid_langs.append(l)
    
    if not valid_langs:
        logger.warning(f"Ningún idioma válido en '{lang}'. Usando español como respaldo.")
        return 'spa'
    
    return '+'.join(valid_langs)

# Preprocesamiento mínimo conservador
def simple_preprocess(image):
    """Mejora básica de imagen sin alterar características clave"""
    img = np.array(image)
    if len(img.shape) == 3:  # Solo si es color
        gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
    else:
        gray = img
    
    # Solo mejora de contraste si es necesario
    if np.mean(gray) < 100 or np.mean(gray) > 200:
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        return clahe.apply(gray)
    return gray

def ProcessOCR(file_storage, fmt, lang='spa'):
    try:
        # Validar idioma
        lang = validate_language(lang)
        logger.info(f"Usando idioma(s): {lang}")
        
        # Cargar imagen
        img = Image.open(file_storage)
        
        # Preprocesamiento mínimo
        processed_img = simple_preprocess(img)
        
        # Convertir a PIL para Tesseract
        img_pil = Image.fromarray(processed_img)
        
        # Configuración mejorada
        tesseract_config = '--oem 3 --psm 6 -c preserve_interword_spaces=1'
        
        # Extraer texto plano
        raw_text = pytesseract.image_to_string(img_pil, lang=lang, config=tesseract_config)
        
        # Extraer datos de layout
        data = pytesseract.image_to_data(
            img_pil, 
            lang=lang, 
            output_type=pytesseract.Output.DICT,
            config=tesseract_config
        )
        
        # Generar documento
        tmp_dir = tempfile.gettempdir()
        filename_base = os.path.splitext(file_storage.filename)[0]
        output_filename = f"resultado_{filename_base}"
        
        if fmt == 'docx':
            path = os.path.join(tmp_dir, f'{output_filename}.docx')
            doc = Document()
            
            # Agrupar texto por bloques y líneas
            current_block = []
            last_block = -1
            last_line = -1
            
            for i in range(len(data['text'])):
                text = data['text'][i].strip()
                if text and int(data['conf'][i]) > 60:  # Filtro de confianza
                    # Nuevo bloque o nueva línea
                    if data['block_num'][i] != last_block or data['line_num'][i] != last_line:
                        if current_block:
                            doc.add_paragraph(' '.join(current_block))
                            current_block = []
                        last_block = data['block_num'][i]
                        last_line = data['line_num'][i]
                    
                    current_block.append(text)
            
            if current_block:
                doc.add_paragraph(' '.join(current_block))
                
            doc.save(path)
            
        elif fmt == 'xlsx':
            path = os.path.join(tmp_dir, f'{output_filename}.xlsx')
            wb = Workbook()
            ws = wb.active
            
            # Organizar texto en filas con agrupación inteligente
            rows = {}
            min_height = 20  # Altura mínima estimada de línea
            
            # Calcular altura promedio de caracteres
            heights = [h for h in data['height'] if h > 0]
            avg_height = sum(heights) / len(heights) if heights else 20
            
            for i in range(len(data['text'])):
                text = data['text'][i].strip()
                if text and int(data['conf'][i]) > 60:
                    y = data['top'][i]
                    
                    # Agrupar por rangos dinámicos basados en altura promedio
                    row_idx = y // max(avg_height, min_height)
                    
                    if row_idx not in rows:
                        rows[row_idx] = []
                    
                    # Mantener información de posición para posible separación columnas
                    rows[row_idx].append({
                        'text': text,
                        'x': data['left'][i],
                        'width': data['width'][i]
                    })
            
            # Escribir en Excel con detección de columnas
            for i, (_, items) in enumerate(sorted(rows.items())):
                # Ordenar elementos por posición X
                sorted_items = sorted(items, key=lambda x: x['x'])
                
                # Intentar detectar columnas
                line_text = []
                current_x = -1
                
                for item in sorted_items:
                    # Si hay un salto significativo en X, podría ser nueva columna
                    if current_x > 0 and (item['x'] - current_x) > avg_height * 2:
                        line_text.append("|")  # Marcador de separación
                    
                    line_text.append(item['text'])
                    current_x = item['x'] + item['width']
                
                # Unir con espacios normales pero mantener separadores de columna
                cell_value = " ".join(line_text).replace(" | ", "|").replace("| ", "|").replace(" |", "|")
                
                # Si hay pipes, separar en columnas
                if '|' in cell_value:
                    columns = cell_value.split('|')
                    for col_idx, col_val in enumerate(columns, 1):
                        ws.cell(row=i+1, column=col_idx, value=col_val.strip())
                else:
                    ws.cell(row=i+1, column=1, value=cell_value)
            
            wb.save(path)
            
        else:  # Texto plano
            path = os.path.join(tmp_dir, f'{output_filename}.txt')
            with open(path, 'w', encoding='utf-8') as f:
                f.write(raw_text)
                
        logger.info(f"Archivo generado: {path}")
        return path
        
    except Exception as e:
        logger.error(f"Error en ProcessOCR: {str(e)}", exc_info=True)
        raise
